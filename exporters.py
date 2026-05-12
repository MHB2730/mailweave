"""MailWeave — PDF and Word export."""

from __future__ import annotations

import html
import os
import zipfile
from datetime import datetime
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from settings import AppSettings
    from email_parser import EmailData

try:
    from reportlab.lib.pagesizes import A4, letter, legal
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        BaseDocTemplate, PageTemplate, Frame, Paragraph, Spacer,
        HRFlowable, KeepTogether, Table, TableStyle,
    )
    from reportlab.platypus.tableofcontents import TableOfContents
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ── Helpers ────────────────────────────────────────────────────────────────────

_DATE_FMTS = {
    'uk':  '%d/%m/%Y %H:%M',
    'us':  '%m/%d/%Y %H:%M',
    'iso': '%Y-%m-%d %H:%M',
}

_PDF_SIZES = {
    'A4':     None,   # filled lazily
    'Letter': None,
    'Legal':  None,
}


def _fmt_long(dt, fallback: str, date_format: str = 'uk') -> str:
    long_fmts = {
        'uk':  '%A, %d %B %Y at %H:%M',
        'us':  '%A, %B %d %Y at %H:%M',
        'iso': '%Y-%m-%d %H:%M',
    }
    fmt = long_fmts.get(date_format, long_fmts['uk'])
    if dt is None:
        return fallback
    try:
        return dt.strftime(fmt)
    except Exception:
        return str(dt)


def _safe(text: str) -> str:
    """Escape text for ReportLab Paragraph XML."""
    return html.escape(str(text or ''), quote=False)


def _choose_body(em: 'EmailData', strip: bool) -> str:
    return em.body_clean if strip else em.body_plain


def _email_summary(em: 'EmailData', strip: bool, limit: int = 80) -> str:
    body = _choose_body(em, strip)
    normalized = ' '.join((body or '').split())
    if not normalized:
        return 'No body preview available.'
    if len(normalized) <= limit:
        return normalized
    return normalized[: limit - 1].rstrip() + '…'


def _index_label(em: 'EmailData', idx: int, strip: bool) -> str:
    subject = (em.subject or '(No Subject)').strip()
    if len(subject) > 90:
        subject = subject[:89].rstrip() + '…'
    summary = _email_summary(em, strip)
    return f'Email {idx + 1}: {subject} — {summary}'


def _email_date_bounds(emails: list['EmailData']) -> tuple[str, str]:
    dated = [email.date for email in emails if email.date is not None]
    if not dated:
        return 'Unknown', 'Unknown'
    return min(dated).isoformat(sep=' ', timespec='minutes'), max(dated).isoformat(sep=' ', timespec='minutes')


def verify_export_file(filepath: str) -> None:
    if not filepath or not os.path.isfile(filepath):
        raise RuntimeError('The export did not produce a file on disk.')
    if os.path.getsize(filepath) <= 0:
        raise RuntimeError('The export file is empty.')
    lower = filepath.lower()
    if lower.endswith('.docx'):
        with zipfile.ZipFile(filepath, 'r') as archive:
            required = {'[Content_Types].xml', 'word/document.xml'}
            names = set(archive.namelist())
            if not required.issubset(names):
                raise RuntimeError('The Word export is incomplete or corrupt.')
    elif lower.endswith('.pdf'):
        with open(filepath, 'rb') as handle:
            header = handle.read(5)
            if header != b'%PDF-':
                raise RuntimeError('The PDF export is incomplete or corrupt.')


def export_annexure_files(base_filepath: str, emails: list['EmailData']) -> str:
    """
    Saves all attachments to a subfolder next to the exported document.
    Returns the name of the folder created.
    """
    import shutil
    from pathlib import Path
    
    base = Path(base_filepath)
    folder_name = f"{base.stem}_Annexures"
    target_dir = base.parent / folder_name
    
    if target_dir.exists():
        shutil.rmtree(target_dir)
    target_dir.mkdir(parents=True, exist_ok=True)
    
    for email in emails:
        for att in email.attachments:
            if not att.content:
                continue
            
            # Clean filename
            label = att.annexure_id or "Attachment"
            clean_name = f"{label} - {att.filename}"
            # Remove invalid chars
            for char in '<>:"/\\|?*':
                clean_name = clean_name.replace(char, '_')
                
            dest = target_dir / clean_name
            dest.write_bytes(att.content)
            
    return folder_name


class _IndexDocTemplate(BaseDocTemplate):
    def __init__(self, filename: str, **kwargs):
        bates_label = kwargs.pop('bates_label', '')
        super().__init__(filename, **kwargs)
        # Inset frame so the footer text has room beneath it.
        frame = Frame(
            self.leftMargin,
            self.bottomMargin + 0.6 * cm,
            self.width,
            self.height - 0.6 * cm,
            id='normal',
        )

        def _draw_footer(canvas, doc):
            canvas.saveState()
            canvas.setFont('Helvetica', 7.5)
            canvas.setFillColor(colors.HexColor('#6B7280'))
            page_w = doc.pagesize[0]
            canvas.drawString(
                doc.leftMargin,
                0.8 * cm,
                bates_label or 'MailWeave',
            )
            canvas.drawRightString(
                page_w - doc.rightMargin,
                0.8 * cm,
                f'Page {doc.page}',
            )
            canvas.restoreState()

        self.addPageTemplates([
            PageTemplate(id='normal', frames=[frame], onPage=_draw_footer),
        ])

    def afterFlowable(self, flowable):
        toc_text = getattr(flowable, '_mw_toc_text', None)
        toc_level = getattr(flowable, '_mw_toc_level', None)
        if toc_text is not None and toc_level is not None:
            self.notify('TOCEntry', (toc_level, toc_text, self.page))


# ── PDF ────────────────────────────────────────────────────────────────────────

def build_pdf(filepath: str,
              emails: list['EmailData'],
              settings: 'AppSettings',
              progress_cb=None,
              cancel_event=None) -> None:
    if not HAS_PDF:
        raise RuntimeError(
            'reportlab is not installed.  Run: pip install reportlab')

    size_map = {'A4': A4, 'Letter': letter, 'Legal': legal}
    page_size = size_map.get(settings.pdf_page_size, A4)
    strip      = settings.strip_quotes
    author     = settings.document_author or 'MailWeave'
    cover      = settings.include_cover_page
    total      = len(emails)

    bates_label = f'{author} — Email Correspondence — {total} email{"s" if total != 1 else ""}'
    doc = _IndexDocTemplate(
        filepath, pagesize=page_size,
        leftMargin=2.2*cm, rightMargin=2.2*cm,
        topMargin=2.2*cm, bottomMargin=2.2*cm,
        title='Email Correspondence', author=author,
        bates_label=bates_label,
    )

    # Colours
    BLACK  = colors.HexColor('#1A2840')
    ACCENT = colors.HexColor('#4A90D9')
    GREY   = colors.HexColor('#6B7280')
    LGREY  = colors.HexColor('#E5E7EB')
    BODY   = colors.HexColor('#2C3040')

    font_pts = {'small': 8.5, 'medium': 9.5, 'large': 11.0}
    body_fs = font_pts.get(settings.font_size, 9.5)

    title_sty = ParagraphStyle(
        'MW_T', fontName='Helvetica-Bold',
        fontSize=22, textColor=BLACK, spaceAfter=4)
    meta_sty = ParagraphStyle(
        'MW_M', fontName='Helvetica',
        fontSize=9, textColor=GREY, spaceAfter=18)
    index_title_sty = ParagraphStyle(
        'MW_IT', fontName='Helvetica-Bold',
        fontSize=16, textColor=BLACK, spaceAfter=10)
    toc_entry_sty = ParagraphStyle(
        'MW_TOC',
        fontName='Helvetica',
        fontSize=8.5,
        textColor=BODY,
        leading=11,
        leftIndent=0,
        firstLineIndent=0,
    )
    num_sty = ParagraphStyle(
        'MW_N', fontName='Helvetica-Bold',
        fontSize=8, textColor=ACCENT, spaceAfter=3)
    subj_sty = ParagraphStyle(
        'MW_S', fontName='Helvetica-Bold',
        fontSize=13, textColor=BLACK, spaceAfter=5, leading=16)
    hdr_lbl = ParagraphStyle(
        'MW_HL', fontName='Helvetica-Bold', fontSize=8.5, textColor=GREY)
    hdr_val = ParagraphStyle(
        'MW_HV', fontName='Helvetica', fontSize=8.5, textColor=BLACK)
    body_sty = ParagraphStyle(
        'MW_B', fontName='Helvetica',
        fontSize=body_fs, textColor=BODY, leading=body_fs * 1.45,
        spaceAfter=4)

    story = []

    if cover:
        story.append(Paragraph('Email Correspondence', title_sty))
        start_date, end_date = _email_date_bounds(emails)
        story.append(Paragraph(
            f'Compiled by {_safe(author)} &nbsp;&bull;&nbsp; '
            f'{datetime.now().strftime("%d %B %Y")} &nbsp;&bull;&nbsp; '
            f'{len(emails)} email{"s" if len(emails) != 1 else ""}',
            meta_sty,
        ))
        manifest_rows = [
            [Paragraph('<b>Date range:</b>', hdr_lbl), Paragraph(_safe(f'{start_date} to {end_date}'), hdr_val)],
            [Paragraph('<b>Page size:</b>', hdr_lbl), Paragraph(_safe(settings.pdf_page_size), hdr_val)],
            [Paragraph('<b>Quote stripping:</b>', hdr_lbl), Paragraph(_safe('Enabled' if strip else 'Disabled'), hdr_val)],
        ]
        manifest = Table(manifest_rows, colWidths=[2.5*cm, None])
        manifest.setStyle(TableStyle([
            ('VALIGN',        (0, 0), (-1, -1), 'TOP'),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
            ('TOPPADDING',    (0, 0), (-1, -1), 2),
            ('LEFTPADDING',   (0, 0), (-1, -1), 5),
            ('RIGHTPADDING',  (0, 0), (-1, -1), 5),
            ('BACKGROUND',    (0, 0), (-1, -1), colors.HexColor('#F8FAFC')),
        ]))
        story.append(manifest)
        story.append(Spacer(1, 10))
        story.append(HRFlowable(width='100%', thickness=2,
                                color=ACCENT, spaceAfter=20))

    if settings.include_email_index and emails:
        toc = TableOfContents()
        toc.levelStyles = [toc_entry_sty]
        story.append(Paragraph('Email Index', index_title_sty))
        story.append(Paragraph(
            'Each entry lists the email subject, a brief body preview, and the page where that email begins.',
            meta_sty,
        ))
        story.append(toc)
        story.append(Spacer(1, 16))
        story.append(HRFlowable(width='100%', thickness=0.5,
                                color=LGREY, spaceAfter=14))

    from annexures import annexure_anchor_id

    for idx, em in enumerate(emails):
        if cancel_event is not None and cancel_event.is_set():
            raise RuntimeError('Export cancelled by user.')
        if progress_cb is not None:
            try:
                progress_cb(idx, total, 'PDF')
            except Exception:
                # Builder treats progress_cb raising as the cancel signal.
                raise

        body_text  = _choose_body(em, strip)
        date_disp  = _fmt_long(em.date, em.date_str, settings.date_format)
        toc_text = _safe(_index_label(em, idx, strip))
        heading = Paragraph(toc_text, subj_sty)
        heading._mw_toc_text = toc_text
        heading._mw_toc_level = 0

        block = [
            Paragraph(f'EMAIL {idx + 1} of {len(emails)}', num_sty),
            heading,
        ]

        rows = [
            [Paragraph('<b>From:</b>',    hdr_lbl), Paragraph(_safe(em.sender),     hdr_val)],
            [Paragraph('<b>To:</b>',      hdr_lbl), Paragraph(_safe(em.recipients), hdr_val)],
            [Paragraph('<b>Date:</b>',    hdr_lbl), Paragraph(_safe(date_disp),     hdr_val)],
        ]
        tbl = Table(rows, colWidths=[1.6*cm, None])
        tbl.setStyle(TableStyle([
            ('VALIGN',        (0, 0), (-1, -1), 'TOP'),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
            ('TOPPADDING',    (0, 0), (-1, -1), 2),
            ('LEFTPADDING',   (0, 0), (-1, -1), 5),
            ('RIGHTPADDING',  (0, 0), (-1, -1), 5),
            ('BACKGROUND',    (0, 0), (-1, -1), colors.HexColor('#F3F4F6')),
        ]))
        block.append(tbl)
        block.append(Spacer(1, 10))
        story.append(KeepTogether(block))

        # Body paragraphs
        buf: list[str] = []
        for line in body_text.splitlines():
            s = line.strip()
            if s:
                buf.append(_safe(s))
            else:
                if buf:
                    story.append(Paragraph(' '.join(buf), body_sty))
                    buf = []
                else:
                    story.append(Spacer(1, 4))
        if buf:
            story.append(Paragraph(' '.join(buf), body_sty))

        # Annexure markers after body — anchored so cross-references in the
        # text can link to them (ReportLab uses <a name="…"/>).
        if em.attachments:
            story.append(Spacer(1, 6))
            for att in em.attachments:
                label = att.annexure_id or "Attachment"
                anchor = annexure_anchor_id(att)
                story.append(Paragraph(
                    f'<a name="{anchor}"/><b>&rarr; {label}:</b> '
                    f'{_safe(att.filename)} ({att.size // 1024} KB)',
                    hdr_val
                ))

        if idx < len(emails) - 1:
            story.append(Spacer(1, 14))
            story.append(HRFlowable(width='100%', thickness=0.5,
                                    color=LGREY, spaceAfter=14))

    doc.multiBuild(story)
    try:
        export_annexure_files(filepath, emails)
    except Exception:
        import logging
        logging.getLogger('MailWeave').exception('annexure-export-failed')


def _set_update_fields_on_open(doc: Document):
    settings = doc.settings.element
    update = settings.find(qn('w:updateFields'))
    if update is None:
        update = OxmlElement('w:updateFields')
        settings.append(update)
    update.set(qn('w:val'), 'true')


def _insert_toc(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-1" \\h \\z \\u'

    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')

    placeholder = OxmlElement('w:t')
    placeholder.text = 'Right-click and update field to refresh page numbers.'

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)


# ── Word ───────────────────────────────────────────────────────────────────────

def build_docx(filepath: str,
               emails: list['EmailData'],
               settings: 'AppSettings',
               progress_cb=None,
               cancel_event=None) -> None:
    if not HAS_DOCX:
        raise RuntimeError(
            'python-docx is not installed.  Run: pip install python-docx')

    strip  = settings.strip_quotes
    author = settings.document_author or 'MailWeave'
    cover  = settings.include_cover_page

    font_pts = {'small': 8.5, 'medium': 10.0, 'large': 12.0}
    body_fs  = font_pts.get(settings.font_size, 10.0)

    doc = Document()
    _set_update_fields_on_open(doc)
    for section in doc.sections:
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    def _r(para, text, bold=False, pts=10.0,
           rgb=(0x1A, 0x28, 0x40)):
        r = para.add_run(str(text))
        r.bold = bold
        r.font.size = Pt(pts)
        r.font.color.rgb = RGBColor(*rgb)
        return r

    if cover:
        p = doc.add_paragraph()
        _r(p, 'Email Correspondence', bold=True, pts=22)
        start_date, end_date = _email_date_bounds(emails)

        p = doc.add_paragraph()
        _r(p,
           f'Compiled by {author}  \u2022  '
           f'{datetime.now().strftime("%d %B %Y")}  \u2022  '
           f'{len(emails)} email{"s" if len(emails) != 1 else ""}',
           pts=9, rgb=(0x6B, 0x72, 0x80))

        for label, value in (
            ('Date range: ', f'{start_date} to {end_date}'),
            ('Page size: ', settings.pdf_page_size),
            ('Quote stripping: ', 'Enabled' if strip else 'Disabled'),
        ):
            p = doc.add_paragraph()
            _r(p, label, bold=True, pts=9, rgb=(0x6B, 0x72, 0x80))
            _r(p, value, pts=9)

        doc.add_paragraph()

    if settings.include_email_index and emails:
        p = doc.add_paragraph()
        p.style = doc.styles['Heading 1']
        _r(p, 'Email Index', bold=True, pts=16)

        p = doc.add_paragraph()
        _r(
            p,
            'This index lists each email with a brief description and its starting page.',
            pts=9,
            rgb=(0x6B, 0x72, 0x80),
        )

        p = doc.add_paragraph()
        _insert_toc(p)
        doc.add_paragraph()

    total_em = len(emails)
    for idx, em in enumerate(emails):
        if cancel_event is not None and cancel_event.is_set():
            raise RuntimeError('Export cancelled by user.')
        if progress_cb is not None:
            progress_cb(idx, total_em, 'Word')

        body_text = _choose_body(em, strip)
        date_disp = _fmt_long(em.date, em.date_str, settings.date_format)

        p = doc.add_paragraph()
        _r(p, f'EMAIL {idx + 1} of {total_em}',
           bold=True, pts=8, rgb=(0x4A, 0x90, 0xD9))

        p = doc.add_paragraph()
        p.style = doc.styles['Heading 1']
        _r(p, _index_label(em, idx, strip), bold=True, pts=13)
        p.paragraph_format.space_after = Pt(4)

        for label, value in (
            ('From:   ', em.sender),
            ('To:       ', em.recipients),
            ('Date:    ', date_disp),
        ):
            p = doc.add_paragraph()
            _r(p, label, bold=True, pts=9, rgb=(0x6B, 0x72, 0x80))
            _r(p, value, pts=9)
            p.paragraph_format.space_after = Pt(2)

        doc.add_paragraph()

        for chunk in body_text.split('\n\n'):
            chunk = chunk.strip()
            if chunk:
                p = doc.add_paragraph()
                _r(p, chunk.replace('\n', ' '), pts=body_fs)
                p.paragraph_format.space_after = Pt(5)

        # Annexure markers
        if em.attachments:
            for att in em.attachments:
                label = att.annexure_id or "Attachment"
                p = doc.add_paragraph()
                _r(p, f"\u2192 {label}: ", bold=True, pts=9, rgb=(0x4A, 0x90, 0xD9))
                _r(p, f"{att.filename} ({att.size // 1024} KB)", pts=9)
                p.paragraph_format.space_after = Pt(2)

        if idx < len(emails) - 1:
            doc.add_paragraph()
            p = doc.add_paragraph()
            _r(p, '\u2500' * 80, pts=5, rgb=(0xCC, 0xCC, 0xCC))
            doc.add_paragraph()

    doc.save(filepath)
    try:
        export_annexure_files(filepath, emails)
    except Exception:
        import logging
        logging.getLogger('MailWeave').exception('annexure-export-failed')
